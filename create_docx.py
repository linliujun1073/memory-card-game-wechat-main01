import os
import sys

report_path = r"c:\Users\Administrator\Desktop\memory-card-game-wechat-main\实验报告.md"
docx_path = r"c:\Users\Administrator\Desktop\memory-card-game-wechat-main\实验报告.docx"

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    with open(report_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
        elif line.startswith('- '):
            doc.add_paragraph(line, style='List Bullet')
        elif line.startswith('| '):
            doc.add_paragraph(line)
        elif line.strip() == '':
            doc.add_paragraph()
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Word document created: {docx_path}")

except ImportError:
    print("python-docx not installed. Installing...")
    os.system("pip install python-docx")
    print("Please run this script again.")
except Exception as e:
    print(f"Error: {e}")